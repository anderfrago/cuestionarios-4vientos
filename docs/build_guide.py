from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).parent / "Guia_instalacion_y_despliegue.docx"
NAVY = RGBColor(23, 50, 77); BLUE = RGBColor(11, 93, 120); TEAL = RGBColor(42, 166, 161)
GRAY = RGBColor(92, 107, 120); LIGHT = "E8EEF5"; BORDER = "CCD7E0"


def font(run, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri"); run.font.size = Pt(size)
    run.bold = bold; run.italic = italic; run.font.color.rgb = color


def shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr(); tcMar = tc.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tc.append(tcMar)
    for side, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = OxmlElement(f"w:{side}"); node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa"); tcMar.append(node)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    tPr=t._tbl.tblPr; tblW=tPr.first_child_found_in("w:tblW"); tblW.set(qn("w:w"),"9360"); tblW.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tPr.append(ind)
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); shade(c,LIGHT); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),10,bold=True,color=BLUE)
    trPr = t.rows[0]._tr.get_or_add_trPr(); repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); trPr.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        for i,(value,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(value)),9.5,color=NAVY)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t


def para(doc, text="", bold=False, color=NAVY, size=11, italic=False, align=None, after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.25
    if align is not None: p.alignment=align
    font(p.add_run(text),size,bold,color,italic); return p


def bullets(doc, items, numbered=False):
    for item in items:
        p=doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.188)
        p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25
        font(p.add_run(item),11)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(1); sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=NAVY
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,before,after,color in (("Title",30,0,8,NAVY),("Subtitle",15,0,20,BLUE),("Heading 1",16,18,10,BLUE),("Heading 2",13,14,7,BLUE),("Heading 3",12,10,5,NAVY)):
    s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=color; s.font.bold=name.startswith("Heading")
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run("AUTOPERCEPCIÓN · GUÍA TÉCNICA"),9,bold=True,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Cuatrovientos · Documento de despliegue y operación"),9,color=GRAY)

para(doc,"GUÍA TÉCNICA",bold=True,color=TEAL,size=11,align=WD_ALIGN_PARAGRAPH.CENTER,after=18)
para(doc,"Autopercepción del alumnado",bold=True,size=30,align=WD_ALIGN_PARAGRAPH.CENTER,after=8)
para(doc,"Instalación, configuración y despliegue en PythonAnywhere",color=BLUE,size=15,align=WD_ALIGN_PARAGRAPH.CENTER,after=28)
para(doc,"Flask · Angular 21 · Signals · Bootstrap · Sass · SQLite",bold=True,color=GRAY,size=10.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=80)
table(doc,["Versión","Destino","Fecha"],[["1.0","Equipo de administración y mantenimiento","1 de septiembre de 2026"]],[1.0,3.6,1.9])
para(doc,"Objetivo",bold=True,color=BLUE,size=12,after=4)
para(doc,"Poner en producción una aplicación que permita al alumnado completar cuestionarios de primero y segundo, conservar su evolución y ofrecer análisis restringidos a cada tutoría.",after=16)
para(doc,"Privacidad",bold=True,color=BLUE,size=12,after=4)
para(doc,"No se importan las respuestas históricas de los archivos de referencia. La política institucional está disponible en https://cuatrovientos.org/rgpd/.",italic=True,color=GRAY)
doc.add_page_break()

heading(doc,"1. Alcance funcional",1)
para(doc,"La solución separa con claridad tres perfiles y aplica el permiso tanto en la interfaz como en cada endpoint de la API.")
table(doc,["Perfil","Puede hacer","Restricción principal"],[["Alumno/a","Registrarse, unirse mediante código, responder varias veces y ver sus resultados","Solo sus cursos e intentos"],["Tutor/a","Consultar medias, niveles, evolución, detalle individual y exportación","Solo cursos asignados"],["Administrador/a","CRUD de cursos, usuarios, aspectos, ítems, rangos y mensajes","Acceso global"]],[1.1,3.2,2.2])
heading(doc,"1.1 Criterios de puntuación",2)
bullets(doc,["Escala: Nunca = 1, A veces = 2, En la mayoría de las veces = 3, Siempre = 4.","Los ítems formulados en negativo usan puntuación inversa: 5 menos la respuesta.","Rangos iniciales editables: 1,00–1,99 Incipiente; 2,00–2,99 En desarrollo; 3,00–4,00 Generado.","Cada aspecto dispone de mensajes editables por nivel y el intento genera una frase final de ánimo."])
heading(doc,"1.2 Cuestionarios iniciales",2)
para(doc,"Los datos iniciales se normalizaron a partir del XLSX de primero y del ODT de segundo. Se agrupan en bienestar emocional, habilidades personales/sociales, aprendizaje, hábitos y necesidades percibidas. Administración puede ajustarlos sin modificar código.")

heading(doc,"2. Arquitectura del repositorio",1)
table(doc,["Ruta","Contenido"],[["backend/","Aplicación Flask, modelos, API, autenticación, permisos y datos iniciales"],["frontend/","Angular 21 standalone, servicios con signals, Bootstrap y Sass"],["tests/","Pruebas de integración del backend"],["docs/","Esta guía operativa"],["wsgi.py","Punto de entrada para PythonAnywhere"],[".env.example","Variables necesarias sin secretos"]],[1.6,4.9])
heading(doc,"2.1 Modelo de datos",2)
bullets(doc,["User: identidad, rol, verificación y estado.","Course y Enrollment: curso académico, nivel, tutor, código de invitación y alumnado.","Aspect e Item: cuestionarios configurables con rangos, mensajes y puntuación inversa.","Attempt y Answer: historial inmutable de cada envío y sus respuestas."],numbered=False)

heading(doc,"3. Instalación local",1)
para(doc,"Requisitos: Python 3.10 o superior, Node.js 20 o superior, npm y Git.")
bullets(doc,["Clonar el repositorio: git clone <URL_DEL_REPOSITORIO> && cd autopercepcion","Crear entorno: python -m venv .venv","Activarlo en Linux/macOS: source .venv/bin/activate","Instalar: pip install -r requirements.txt pytest","Copiar configuración: cp .env.example .env","Crear esquema: flask --app backend.app db upgrade","Cargar cuestionarios: flask --app backend.app seed-data","Arrancar API: flask --app backend.app run --debug","En otra terminal: cd frontend && npm install && npm start"],numbered=True)
para(doc,"Angular usa proxy local para enviar /api a http://127.0.0.1:5000. La interfaz queda en http://localhost:4200.",italic=True,color=GRAY)
heading(doc,"3.1 Verificación",2)
bullets(doc,["Backend: .venv/bin/python -m pytest -q","Frontend: cd frontend && npm run build","Salud: GET http://127.0.0.1:5000/api/health debe devolver {status: ok}."])

heading(doc,"4. Variables de entorno",1)
table(doc,["Variable","Uso","Producción"],[["SECRET_KEY / JWT_SECRET_KEY","Sesiones OAuth y tokens JWT","Dos cadenas aleatorias de 32+ bytes"],["DATABASE_URL","Base de datos","sqlite:////home/<usuario>/autopercepcion/instance/autopercepcion.db"],["AUTO_CREATE_DB","Creación automática para demos","false; usar migraciones"],["FRONTEND_URL","Origen CORS y retorno OAuth","https://<usuario>.pythonanywhere.com"],["ADMIN_EMAILS","Administradores separados por comas","Revisar antes del primer acceso"],["GOOGLE_CLIENT_ID / SECRET","OAuth 2.0","Desde Google Cloud"],["MAIL_*","Verificación por correo","SMTP con contraseña de aplicación"]],[2.0,2.2,2.3])
para(doc,"Nunca subir .env al repositorio. El archivo ya está excluido mediante .gitignore.",bold=True,color=BLUE)

heading(doc,"5. Configurar acceso con Google",1)
bullets(doc,["Crear un proyecto en Google Cloud Console y configurar la pantalla de consentimiento OAuth.","Crear credenciales de tipo Aplicación web.","Añadir como origen autorizado la URL pública de PythonAnywhere.","Añadir URI de redirección: https://<usuario>.pythonanywhere.com/api/auth/google/callback.","Copiar identificador y secreto en .env.","Publicar la aplicación OAuth o añadir cuentas de prueba mientras permanezca en modo prueba."],numbered=True)
para(doc,"Las cuentas @educacion.navarra.es funcionan si están administradas por Google Workspace y la política del dominio permite autorizar esta aplicación.",italic=True,color=GRAY)
heading(doc,"5.1 Registro por correo",2)
para(doc,"Configurar MAIL_SERVER, MAIL_PORT, MAIL_USE_TLS, MAIL_USERNAME, MAIL_PASSWORD y MAIL_DEFAULT_SENDER. Para Gmail se recomienda activar verificación en dos pasos y crear una contraseña de aplicación; no usar la contraseña normal de la cuenta.")

heading(doc,"6. Preparar GitHub",1)
bullets(doc,["Crear un repositorio vacío en GitHub sin README adicional.","Ejecutar git init, git add . y git commit -m \"Aplicación inicial\".","Añadir remoto: git remote add origin <URL_DEL_REPOSITORIO>.","Publicar: git branch -M main && git push -u origin main.","Comprobar que .env, instance/, node_modules/ y bases .db no aparecen en GitHub."],numbered=True)

heading(doc,"7. Despliegue en PythonAnywhere",1)
para(doc,"Este procedimiento contempla una única URL pública. Angular se compila y Flask sirve el resultado estático en producción; si se mantiene separado, habría que contratar o configurar un segundo alojamiento para el frontend.")
bullets(doc,["Crear cuenta y abrir una consola Bash.","Clonar: git clone <URL_DEL_REPOSITORIO> ~/autopercepcion.","Crear virtualenv: python3.10 -m venv ~/.virtualenvs/autopercepcion.","Instalar: ~/.virtualenvs/autopercepcion/bin/pip install -r ~/autopercepcion/requirements.txt.","Compilar Angular localmente o en un entorno con Node: cd frontend && npm ci && npm run build. Subir frontend/dist/autopercepcion/browser al repositorio de despliegue o copiarlo al servidor.","Crear ~/autopercepcion/.env con valores reales y AUTO_CREATE_DB=false.","Ejecutar flask --app backend.app db upgrade y flask --app backend.app seed-data.","En Web > Add a new web app, elegir Manual configuration y la versión de Python disponible.","Indicar virtualenv: /home/<usuario>/.virtualenvs/autopercepcion.","Editar el archivo WSGI según el bloque siguiente y pulsar Reload."],numbered=True)
heading(doc,"7.1 Archivo WSGI",2)
para(doc,"Añadir al WSGI de PythonAnywhere:",bold=True)
para(doc,"import sys\nfrom pathlib import Path\nproject = Path('/home/<usuario>/autopercepcion')\nsys.path.insert(0, str(project))\nfrom dotenv import load_dotenv\nload_dotenv(project / '.env')\nfrom backend import create_app\napplication = create_app()",size=9.5,color=RGBColor(30,45,55))
heading(doc,"7.2 Archivos estáticos",2)
para(doc,"En Web > Static files, mapear /assets/ al directorio frontend/dist/autopercepcion/browser/assets/. Para una SPA con rutas cliente, es preferible que Flask entregue index.html como fallback; si se despliega el frontend por separado, establecer FRONTEND_URL a su dominio.")

heading(doc,"8. Primera puesta en marcha",1)
bullets(doc,["Acceder con uno de los correos incluidos en ADMIN_EMAILS; la cuenta recibirá rol admin automáticamente.","Crear o promover cuentas de tutoría desde Administración > Usuarios.","Crear cursos con nivel primero o segundo y año académico.","Asignar un tutor a cada curso.","Compartir el código de invitación únicamente con el alumnado del grupo.","Revisar aspectos, ítems, rangos y mensajes antes de abrir el cuestionario.","Realizar un envío de prueba y comprobar panel, detalle y exportación."],numbered=True)
heading(doc,"8.1 Actualizaciones",2)
bullets(doc,["En consola: cd ~/autopercepcion && git pull.","Actualizar paquetes si cambió requirements.txt.","Reconstruir Angular si cambió frontend/.","Aplicar migraciones cuando existan: flask --app backend.app db upgrade.","Recargar la web desde el panel de PythonAnywhere."])

heading(doc,"9. Seguridad y protección de datos",1)
table(doc,["Control","Aplicación"],[["Mínimo privilegio","Validación de rol y pertenencia al curso en backend"],["Contraseñas","Hash seguro de Werkzeug; nunca se almacenan en claro"],["Verificación","Correo de confirmación o identidad Google"],["Datos históricos","Intentos conservados para evolución; tutoría no puede editar"],["Transporte","HTTPS obligatorio en producción"],["Privacidad","Enlace institucional RGPD visible; no se importan respuestas de ejemplo"]],[1.8,4.7])
heading(doc,"9.1 Recomendaciones antes de producción",2)
bullets(doc,["Definir plazo de conservación y procedimiento de borrado o anonimización.","Documentar base jurídica, responsables y encargados del tratamiento.","Limitar el acceso administrativo y revisar ADMIN_EMAILS periódicamente.","Crear copias de seguridad cifradas de SQLite y probar su restauración.","No usar los resultados como diagnóstico clínico. Establecer un protocolo humano ante respuestas de riesgo."])
para(doc,"Atención: los ítems relacionados con autolesión requieren un protocolo claro de salvaguarda del centro. La aplicación muestra resultados, pero no sustituye una intervención profesional.",bold=True,color=RGBColor(155,28,28))

heading(doc,"10. Copias de seguridad y mantenimiento",1)
bullets(doc,["Detener escrituras o usar la herramienta de backup de SQLite antes de copiar la base.","Conservar varias generaciones fuera del servidor y registrar fecha de cada copia.","Verificar mensualmente acceso, correo, OAuth, exportación y restauración.","Revisar actualizaciones de dependencias y ejecutar pruebas antes de publicar.","Monitorizar el espacio disponible: SQLite, logs y copias pueden agotarlo."],numbered=True)
heading(doc,"10.1 Lista de aceptación",2)
bullets(doc,["Registro por correo envía enlace y permite verificar.","Google OAuth retorna a la URL pública.","Un código inválido no incorpora a ningún curso.","Un tutor no puede abrir otro curso cambiando la URL.","Todos los intentos aparecen en orden temporal.","Los rangos y mensajes editados se reflejan en resultados nuevos.","La exportación contiene solo alumnado del curso autorizado.","La política RGPD abre correctamente."])

heading(doc,"11. Resolución de problemas",1)
table(doc,["Síntoma","Comprobación"],[["Error 500 al iniciar","Revisar Error log, ruta del virtualenv, .env y permisos de instance/"],["Google redirect_uri_mismatch","La URI debe coincidir exactamente, incluido https y /callback"],["No llega el correo","Revisar contraseña de aplicación, TLS, remitente y Spam"],["CORS bloqueado","FRONTEND_URL debe coincidir exactamente con el origen"],["Base no abre","Usar ruta absoluta y dar escritura al directorio instance"],["Angular muestra 404 al recargar","Configurar fallback SPA o servir index.html desde Flask"],["Tutor ve acceso denegado","Comprobar rol activo y tutor_id del curso"]],[2.1,4.4])

heading(doc,"12. Referencia operativa",1)
para(doc,"Comandos principales",bold=True,color=BLUE,size=12)
table(doc,["Tarea","Comando"],[["Pruebas backend",".venv/bin/python -m pytest -q"],["Compilar frontend","cd frontend && npm run build"],["Servidor local API","flask --app backend.app run --debug"],["Inicializar migraciones","flask --app backend.app db init"],["Crear migración","flask --app backend.app db migrate -m \"descripción\""],["Aplicar migración","flask --app backend.app db upgrade"]],[2.1,4.4])
para(doc,"Fuentes funcionales: Cuestionario inicial-Tutoría (respuestas).xlsx y autopercecionsegundos.odt. No se incorporan sus respuestas históricas.",italic=True,color=GRAY,size=9.5)
para(doc,"Fin de la guía · Revisar las URLs y marcadores <usuario> / <URL_DEL_REPOSITORIO> antes de ejecutar el despliegue.",bold=True,color=TEAL,align=WD_ALIGN_PARAGRAPH.CENTER,after=0)

doc.core_properties.title="Autopercepción del alumnado - Guía de instalación y despliegue"
doc.core_properties.subject="Flask, Angular 21, SQLite y PythonAnywhere"
doc.core_properties.author="Cuatrovientos"
OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)
